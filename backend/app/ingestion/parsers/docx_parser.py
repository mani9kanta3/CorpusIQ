"""
DOCX Parser using python-docx
=============================

Why python-docx?
----------------
1. OFFICIAL-ISH: Most widely used library for .docx files
2. STRUCTURE: Preserves document structure (paragraphs, tables, headers)
3. SIMPLE API: Easy to use compared to parsing XML manually

What is a DOCX file?
--------------------
DOCX is Microsoft's Open XML format. A .docx file is actually a ZIP archive:

    document.docx (rename to .zip and extract)
    ├── [Content_Types].xml
    ├── _rels/
    ├── docProps/
    │   ├── app.xml        # Application metadata
    │   └── core.xml       # Author, title, dates
    └── word/
        ├── document.xml   # THE ACTUAL CONTENT
        ├── styles.xml     # Formatting styles
        └── ...

python-docx reads this structure and gives us a clean Python API.

Installation: pip install python-docx
Import as: from docx import Document (NOT import docx)
"""

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pathlib import Path
from datetime import datetime
from typing import Optional

from .base import BaseParser, DocumentMetadata, ParsedDocument


class DOCXParser(BaseParser):
    """
    Parser for Microsoft Word documents (.docx).
    
    Note on "pages":
    ----------------
    Unlike PDFs, Word documents don't have fixed pages. Page breaks depend on:
    - Paper size settings
    - Margins
    - Font sizes
    - The application rendering it
    
    So our 'pages' list actually contains PARAGRAPHS, not pages.
    Each paragraph is a separate item. This is still useful for citations
    because we can reference "paragraph 15" instead of "page 3".
    
    Limitation:
    -----------
    This parser extracts plain text. It does NOT extract:
    - Images (we'll handle in a later module)
    - Complex formatting
    - Headers/footers (we'll add this)
    - Comments/track changes
    
    For now, we focus on getting the text content out.
    
    Usage:
        parser = DOCXParser()
        result = parser.parse(Path("report.docx"))
        print(result.content)
        print(len(result.pages))  # Number of paragraphs
    """
    
    SUPPORTED_EXTENSIONS = [".docx"]
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse a DOCX file and extract all text content.
        
        Extracts:
        - All paragraphs from the main document body
        - Tables (converted to text)
        - Metadata (author, title, dates)
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            ParsedDocument with content, metadata, and paragraphs
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If not a .docx file
            RuntimeError: If parsing fails
        """
        self.validate_file(file_path)
        
        try:
            # Document() opens the DOCX file
            # Unlike fitz, python-docx doesn't need 'with' statement
            # (though it works with 'with' too)
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = self._extract_paragraphs(doc)
            
            # Extract tables and convert to text
            tables_text = self._extract_tables(doc)
            
            # Combine paragraphs into full content
            full_content = "\n\n".join(paragraphs)
            
            # Add tables at the end if any exist
            if tables_text:
                full_content += "\n\n--- Tables ---\n\n" + tables_text
            
            # Extract metadata
            metadata = self._extract_metadata_from_doc(doc, file_path)
            
            return ParsedDocument(
                content=full_content,
                metadata=metadata,
                pages=paragraphs  # Remember: these are paragraphs, not pages
            )
            
        except PackageNotFoundError:
            # This error occurs when the file isn't a valid DOCX
            # (e.g., someone renamed a .txt to .docx)
            raise RuntimeError(
                f"'{file_path}' is not a valid DOCX file. "
                "It may be corrupted or not a real Word document."
            )
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX '{file_path}': {str(e)}")
    
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """
        Extract only metadata from a DOCX file.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            DocumentMetadata object
        """
        self.validate_file(file_path)
        
        try:
            doc = Document(file_path)
            return self._extract_metadata_from_doc(doc, file_path)
        except PackageNotFoundError:
            raise RuntimeError(f"'{file_path}' is not a valid DOCX file.")
        except Exception as e:
            raise RuntimeError(f"Failed to extract metadata from '{file_path}': {str(e)}")
    
    def _extract_paragraphs(self, doc: Document) -> list[str]:
        """
        Extract all paragraphs from the document.
        
        Why filter empty paragraphs?
        ----------------------------
        Word documents often have empty paragraphs for spacing.
        We skip these because:
        1. They add no value to our text content
        2. They'd mess up paragraph numbering for citations
        3. They'd create noise in embeddings
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of non-empty paragraph texts
        """
        paragraphs = []
        
        for para in doc.paragraphs:
            # para.text gets the plain text content
            text = para.text.strip()
            
            # Only include non-empty paragraphs
            if text:
                paragraphs.append(text)
        
        return paragraphs
    
    def _extract_tables(self, doc: Document) -> str:
        """
        Extract tables and convert to readable text format.
        
        Tables in DOCX are separate from paragraphs.
        We convert them to a simple text format:
        
            Row 1: Cell1 | Cell2 | Cell3
            Row 2: Cell1 | Cell2 | Cell3
        
        Later, in Module 3 (Table Extraction), we'll do more sophisticated
        table handling. For now, this ensures we don't lose table data.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            String containing all tables as text
        """
        tables_text = []
        
        for table_idx, table in enumerate(doc.tables):
            table_rows = []
            
            for row in table.rows:
                # Extract text from each cell in the row
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_cells.append(cell_text)
                
                # Join cells with " | " separator
                row_text = " | ".join(row_cells)
                table_rows.append(row_text)
            
            # Combine rows with newlines
            table_text = "\n".join(table_rows)
            tables_text.append(f"Table {table_idx + 1}:\n{table_text}")
        
        # Join all tables
        return "\n\n".join(tables_text)
    
    def _extract_metadata_from_doc(
        self, 
        doc: Document, 
        file_path: Path
    ) -> DocumentMetadata:
        """
        Extract metadata from document core properties.
        
        DOCX metadata is stored in docProps/core.xml inside the ZIP.
        python-docx exposes this via doc.core_properties.
        
        Args:
            doc: python-docx Document object
            file_path: Path to the file (for filename and size)
            
        Returns:
            DocumentMetadata object
        """
        # core_properties contains document metadata
        props = doc.core_properties
        
        # Count paragraphs as our "page count" equivalent
        # This gives users a sense of document length
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        
        return DocumentMetadata(
            filename=file_path.name,
            file_type=self._get_file_extension(file_path),
            file_size_bytes=self._get_file_size(file_path),
            page_count=paragraph_count,  # Actually paragraph count for DOCX
            created_at=props.created,    # Already a datetime object
            modified_at=props.modified,  # Already a datetime object
            author=props.author,
            title=props.title
        )